import csv
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


def convert(csv_path: Path, out_path: Path) -> int:
    rows = []
    with csv_path.open(encoding="utf-8", newline="") as f:
        rows = list(csv.reader(f))

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    title = doc.add_heading("PDFClinic QA Testing Checklist", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    intro = doc.add_paragraph(
        "Comprehensive QA checklist covering homepage, navigation, all 14 PDF tools, "
        "upload/process/download flows, mobile, browser, SEO, accessibility, security, and regression testing."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta = doc.add_paragraph(f"Total test cases: {len(rows) - 1}")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.runs[0].font.size = Pt(10)
    meta.runs[0].font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph()

    headers = rows[0]
    data_rows = rows[1:]

    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False

    widths = [
        Inches(1.0),
        Inches(0.85),
        Inches(0.85),
        Inches(1.35),
        Inches(1.55),
        Inches(0.55),
        Inches(0.75),
        Inches(0.45),
    ]
    for i, w in enumerate(widths):
        for cell in table.columns[i].cells:
            cell.width = w

    hdr_cells = table.rows[0].cells
    for i, text in enumerate(headers):
        cell = hdr_cells[i]
        cell.text = text
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(8)
        shading = OxmlElement("w:shd")
        shading.set(qn("w:fill"), "1F4E79")
        cell._tc.get_or_add_tcPr().append(shading)
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)

    for row in data_rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = val
            for p in cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(7)

    doc.add_page_break()
    doc.add_heading("Priority Legend", level=1)
    doc.add_paragraph("P0 – Critical: Must pass before release")
    doc.add_paragraph("P1 – High: Important functionality and compliance")
    doc.add_paragraph("P2 – Medium: Edge cases and polish")
    doc.add_paragraph("P3 – Low: Nice-to-have")
    doc.add_paragraph()
    doc.add_heading("Status Values", level=1)
    doc.add_paragraph("Not Started | Pass | Fail | Blocked | N/A")

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return len(data_rows)


if __name__ == "__main__":
    src = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(r"C:\Users\HP\Desktop\reports\PDFClinic-QA-Checklist.csv")
    dst = Path(sys.argv[2]) if len(sys.argv) > 2 else src.with_suffix(".docx")
    count = convert(src, dst)
    print(f"Created: {dst}")
    print(f"Rows: {count}")
